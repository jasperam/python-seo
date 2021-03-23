#
# report.py
# @author Neo Lin
# @description SEO risk report
# @created 2020-05-26T14:19:21.122Z+08:00
# @last-modified 2021-03-23T13:33:27.921Z+08:00
#
import os
import docx

from datetime import datetime
from jt.invest.constants import w_db
from jt.utils.fs.utils import Utils as futils
from jt.utils.misc.log import Logger

from seo.constants import root, product_map

log = Logger(__name__)

def make_intent_letter(broker, project):
    """
    make yi xiang letter
    """
    _year = str(datetime.now().year)
    _month = str(datetime.now().month)
    _day = str(datetime.now().day) 
    
    sql = f"""
        select S_INFO_CODE as code, s_info_name as name
        from [dbo].[ASHAREDESCRIPTION] where S_INFO_NAME in ('{"','".join(project)}')
    """
    ret = w_db.read(sql)

    for i in range(len(project)):
        _pro = project[i]
        _broker = broker[i]  
        _doc = docx.Document(os.path.join(root, u'30 合作券商/认购意向函.docx'))
        for _p in _doc.paragraphs:
            if ('【】' in _p.text) or ('broker' in _p.text) or (u'x月y日' in _p.text):
                for _r in _p.runs:
                    if ('【】' in _r.text):
                        _r.text = _r.text.replace('【】', _pro)
                    if ('broker' in _r.text):
                        _r.text = _r.text.replace('broker', _broker)
                    if ('x' in _r.text):
                        _r.text = _r.text.replace(u'x', _month)
                    if ('y' in _r.text):
                        _r.text = _r.text.replace(u'y', _day)
                    if ('z' in _r.text):
                        _r.text = _r.text.replace(u'z', _year)
        _symbol = ret.loc[ret['name']==_pro, 'code'].to_numpy()[0]
        if "*" in _pro[0]:
            _pro=_pro[1:]
        _dir = os.path.join(root, f'00 项目资料\\{_pro}_{_symbol}')
        if not futils.is_dir(_dir):
            futils.make_dir(_dir)

        file_path = os.path.join(_dir, f'{_pro}认购意向函.docx')
        _doc.save(file_path)
        _doc.save(f'C:/Users/jasper/Desktop/x/{_pro}认购意向函.docx')
        log.info(f'make intent file: {file_path}')


def make_commitment_letter(project,product):
    """
    make chen nuo letter
    """
    _year = str(datetime.now().year)
    _month = str(datetime.now().month)
    _day = str(datetime.now().day) 
    
    sql = f"""
        select S_INFO_CODE as code, s_info_name as name, s_info_compname as fullname
        from [dbo].[ASHAREDESCRIPTION] where S_INFO_NAME in ('{"','".join(project)}')
    """
    ret = w_db.read(sql)

    for _pro in project:
        _doc = docx.Document(os.path.join(root, u'30 合作券商/产品承诺函.docx'))
        _symbol = ret.loc[ret['name']==_pro, 'code'].to_numpy()[0]
        _fullname = ret.loc[ret['name']==_pro, 'fullname'].to_numpy()[0]
        for _p in _doc.paragraphs:
            if ('【】' in _p.text) or ('fullname' in _p.text) or ('code' in _p.text) or (u'x月y日' in _p.text) or ('product_name' in _p.text):
                for _r in _p.runs:
                    if ('【】' in _r.text):
                        _r.text = _r.text.replace('【】', _pro)
                    if ('fullname' in _r.text):
                        _r.text = _r.text.replace('fullname', _fullname)
                    if ('full' in _r.text):
                        _r.text = _r.text.replace('full', _fullname)
                    if ('code' in _r.text):
                        _r.text = _r.text.replace('code', _symbol)
                    if ('z' in _r.text):
                        _r.text = _r.text.replace(u'z', _year)
                    if ('x' in _r.text):
                        _r.text = _r.text.replace(u'x', _month)
                    if ('y' in _r.text):
                        _r.text = _r.text.replace(u'y', _day)
                    if ('product_name' in _r.text):
                        _r.text = _r.text.replace('product_name', f'{product_map[product]}')
        _symbol = ret.loc[ret['name']==_pro, 'code'].to_numpy()[0]
        _dir = os.path.join(root, f'00 项目资料\\{_pro}_{_symbol}')
        if not futils.is_dir(_dir):
            futils.make_dir(_dir)

        file_path = os.path.join(_dir, f'{_pro}{product_map[product]}承诺函.docx')
        _doc.save(file_path)
        _doc.save(f'C:/Users/jasper/Desktop/x/{_pro}{product_map[product]}承诺函.docx')
        log.info(f'make intent file: {file_path}')


if __name__ == "__main__":    
    project = [u'新疆天业']
    broker = [u'申万宏源证券承销保荐有限责任公司']
    make_intent_letter(broker, project)



